"""Generate docs/weekly-submissions/Week7_Final_Report.docx.

Throwaway doc-tooling, like build_week6_report.py. Not part of the app, not
covered by the test suite. Mirrors the formatting of the Week 5/6 reports
(Times New Roman, double-spaced, centered title page, bold numbered section
headings, "List Bullet" for bullet lists) so the whole series looks like one
document, and adds three matplotlib figures (rendered to a temp dir, embedded,
then discarded) for the assignment's visual-communication criterion.

    .venv/bin/python scripts/build_week7_report.py
"""

from __future__ import annotations

import tempfile
from pathlib import Path

import matplotlib

matplotlib.use("Agg")

import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "weekly-submissions" / "Week7_Final_Report.docx"

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)

BLUE = "#4C72B0"
ORANGE = "#DD8452"
GRAY = "#8C8C8C"


def _set_run_font(run) -> None:  # type: ignore[no-untyped-def]
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE


def _no_extra_spacing(p) -> None:  # type: ignore[no-untyped-def]
    """Word's default template adds 10pt after every paragraph on top of
    line spacing; zero it so double line-spacing is the only spacing rule,
    matching standard double-spaced manuscript format."""
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 2.0


def add_para(doc: Document, text: str = "", *, center: bool = False, bold: bool = False, italic: bool = False, size: Pt | None = None) -> None:  # type: ignore[no-untyped-def]
    p = doc.add_paragraph()
    _no_extra_spacing(p)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        _set_run_font(run)
        if size is not None:
            run.font.size = size


def add_body(doc: Document, text: str) -> None:  # type: ignore[no-untyped-def]
    """A double-spaced body paragraph, left-aligned (Word's default justify-none)."""
    p = doc.add_paragraph()
    _no_extra_spacing(p)
    run = p.add_run(text)
    _set_run_font(run)


def add_heading(doc: Document, text: str) -> None:  # type: ignore[no-untyped-def]
    add_para(doc, text, bold=True)


def add_subheading(doc: Document, text: str) -> None:  # type: ignore[no-untyped-def]
    p = doc.add_paragraph()
    _no_extra_spacing(p)
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    _set_run_font(run)


def add_bullets(doc: Document, items: list[str]) -> None:  # type: ignore[no-untyped-def]
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        _no_extra_spacing(p)
        run = p.add_run(item)
        _set_run_font(run)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:  # type: ignore[no-untyped-def]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for cell, text in zip(table.rows[0].cells, headers, strict=True):
        _no_extra_spacing(cell.paragraphs[0])
        cell.paragraphs[0].paragraph_format.line_spacing = 1.0
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        _set_run_font(run)
    for row, values in zip(table.rows[1:], rows, strict=True):
        for cell, text in zip(row.cells, values, strict=True):
            _no_extra_spacing(cell.paragraphs[0])
            cell.paragraphs[0].paragraph_format.line_spacing = 1.0
            run = cell.paragraphs[0].add_run(text)
            _set_run_font(run)
    p = doc.add_paragraph()
    _no_extra_spacing(p)


def add_figure(doc: Document, image_path: Path, caption: str) -> None:  # type: ignore[no-untyped-def]
    doc.add_picture(str(image_path), width=Inches(5.0))
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _no_extra_spacing(last)
    add_para(doc, caption, center=True, italic=True, size=Pt(10))


def _style_axes(ax) -> None:  # type: ignore[no-untyped-def]
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)
    ax.grid(axis="y", color="#E0E0E0", linewidth=0.8, zorder=0)
    ax.set_axisbelow(True)


def make_pipeline_diagram(tmpdir: Path) -> Path:
    fig, ax = plt.subplots(figsize=(7.2, 3.6))
    ax.set_xlim(0, 10)
    ax.set_ylim(0, 5)
    ax.axis("off")

    agents = [
        ("Job\nDescription", 0.6),
        ("Evidence\nRetrieval", 2.5),
        ("Resume\nWriter", 4.4),
        ("Trust\nHarness", 6.3),
        ("ATS\nEvaluation", 8.2),
    ]
    box_w, box_h = 1.5, 1.2
    y = 3.0
    for label, x in agents:
        ax.add_patch(
            plt.Rectangle((x, y), box_w, box_h, facecolor="#EDF2FB", edgecolor=BLUE, linewidth=1.5, zorder=2)
        )
        ax.text(x + box_w / 2, y + box_h / 2, label, ha="center", va="center", fontsize=8.5, zorder=3)

    for _, x in agents[:-1]:
        ax.annotate(
            "",
            xy=(x + box_w + 0.35, y + box_h / 2),
            xytext=(x + box_w, y + box_h / 2),
            arrowprops={"arrowstyle": "-|>", "color": "#333333", "lw": 1.3},
        )

    # Quality-loop feedback arrow: Trust/ATS -> Resume Writer (rewrite)
    ax.annotate(
        "",
        xy=(4.4 + box_w / 2, y - 0.15),
        xytext=(6.3 + box_w / 2, y - 0.15),
        arrowprops={"arrowstyle": "-|>", "color": ORANGE, "lw": 1.3, "connectionstyle": "arc3,rad=-0.35"},
    )
    ax.text(5.55, 1.55, "always rewrites N more times;\nkeeps the best-scoring draft", ha="center", fontsize=7.5, color=ORANGE)

    # Candidate Profile agent, feeding the writer, off to the side (cached).
    ax.add_patch(
        plt.Rectangle((4.4, 0.4), box_w, 0.9, facecolor="#FBEEE6", edgecolor=ORANGE, linewidth=1.3, zorder=2)
    )
    ax.text(4.4 + box_w / 2, 0.4 + 0.45, "Candidate Profile\n(cached)", ha="center", va="center", fontsize=7.5, zorder=3)
    ax.annotate(
        "",
        xy=(4.4 + box_w / 2, y),
        xytext=(4.4 + box_w / 2, 0.4 + 0.9),
        arrowprops={"arrowstyle": "-|>", "color": ORANGE, "lw": 1.1},
    )

    ax.text(0.0, 4.65, "One LangGraph StateGraph run", fontsize=9, color="#555555")
    fig.tight_layout()
    path = tmpdir / "fig1_pipeline.png"
    fig.savefig(path, dpi=200)
    plt.close(fig)
    return path


def make_results_chart(tmpdir: Path) -> Path:
    """One combined figure (two panels) instead of two separate ones, to
    keep the report's total figure footprint down."""
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(9.6, 3.2))

    metrics = ["Accuracy", "Macro-F1", "SUPPORTED\nrecall"]
    before = [0.500, 0.489, 0.25]
    after = [0.833, 0.833, 1.00]
    x = range(len(metrics))
    width = 0.32
    ax1.bar([i - width / 2 for i in x], before, width, label="Before", color=GRAY)
    ax1.bar([i + width / 2 for i in x], after, width, label="After", color=BLUE)
    ax1.set_xticks(list(x))
    ax1.set_xticklabels(metrics, fontsize=8)
    ax1.set_ylim(0, 1.05)
    ax1.set_title("Trust Harness calibration", fontsize=9.5)
    _style_axes(ax1)
    ax1.legend(loc="upper left", fontsize=7.5, frameon=False)
    for i, (b, a) in enumerate(zip(before, after, strict=True)):
        ax1.text(i - width / 2, b + 0.03, f"{b:.2f}", ha="center", fontsize=7.5)
        ax1.text(i + width / 2, a + 0.03, f"{a:.2f}", ha="center", fontsize=7.5)

    scenarios = ["1-strong", "2-partial", "3-inflation", "4-wrong-\ndomain"]
    trust = [93, 81, 83, 100]
    ats = [95, 63, 42, 0]
    x2 = range(len(scenarios))
    ax2.bar([i - width / 2 for i in x2], trust, width, label="Trust", color=BLUE)
    ax2.bar([i + width / 2 for i in x2], ats, width, label="ATS", color=ORANGE)
    ax2.axhline(90, color=BLUE, linestyle="--", linewidth=1, alpha=0.6)
    ax2.axhline(85, color=ORANGE, linestyle="--", linewidth=1, alpha=0.6)
    ax2.set_xticks(list(x2))
    ax2.set_xticklabels(scenarios, fontsize=7.5)
    ax2.set_ylim(0, 112)
    ax2.set_title("Four real Bedrock generations", fontsize=9.5)
    _style_axes(ax2)
    ax2.legend(loc="upper right", fontsize=7.5, frameon=False)
    for i, (t, a) in enumerate(zip(trust, ats, strict=True)):
        ax2.text(i - width / 2, t + 2, str(t), ha="center", fontsize=7.5)
        ax2.text(i + width / 2, a + 2, str(a), ha="center", fontsize=7.5)

    fig.tight_layout()
    path = tmpdir / "fig2_results.png"
    fig.savefig(path, dpi=200)
    plt.close(fig)
    return path


def build(tmpdir: Path) -> Document:  # type: ignore[no-untyped-def]
    doc = Document()

    for style_name in ("Normal", "List Bullet"):
        style = doc.styles[style_name]
        style.font.name = FONT_NAME
        style.font.size = FONT_SIZE

    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title page ---------------------------------------------------------
    add_para(doc, "", center=True)
    add_para(doc, "MSAI 699 Capstone Project", center=True)
    add_para(doc, "", center=True)
    add_para(doc, "Instructor: Dr. Gamini Bulumulle", center=True)
    add_para(doc, "", center=True)
    add_para(doc, "Week 7: Final Report", center=True)
    add_para(doc, "", center=True)
    add_para(
        doc,
        "TrustResume: Evidence-Based Resume Generation Using RAG, Multi-Agent AI, "
        "and Trust Verification",
        center=True,
    )
    add_para(doc, "", center=True)
    add_para(doc, "Peng Fei Xu", center=True)
    add_para(doc, "", center=True)
    add_para(doc, "University of the Cumberlands", center=True)
    add_para(doc, "", center=True)
    add_para(doc, "Submitted to the University of the Cumberlands", center=True)
    add_para(doc, "in Partial Fulfillment of the Requirements of the Degree of", center=True)
    add_para(doc, "Master of Science in Artificial Intelligence", center=True)
    add_para(doc, "", center=True)
    add_para(doc, "Aug 11, 2026", center=True)
    doc.add_page_break()

    # --- Abstract ------------------------------------------------------------
    add_para(doc, "Abstract", center=True, bold=True)
    add_body(
        doc,
        "TrustResume checks every claim in a generated resume against a "
        "candidate's own evidence before showing it to anyone. This is a "
        "direct response to the risk that an unconstrained LLM embellishes "
        "a resume. This report covers a from-scratch reimplementation of an "
        "earlier pydantic-ai/Qdrant capstone on LangChain, LangGraph, and "
        "ChromaDB: six agents coordinated by a quality loop that always "
        "rewrites a draft a configurable number of times against a Trust "
        "gate and an ATS gate, regardless of whether an earlier draft "
        "already passed, then ships whichever draft actually scored best. "
        "An offline evaluation harness scores the system itself: hybrid "
        "retrieval reaches recall@8 1.000/MRR 0.938, and the Trust Harness, "
        "after a calibration fix documented here, classifies claims at "
        "accuracy/macro-F1 0.833 (up from 0.500/0.489) with zero dangerous "
        "false-pass errors either way. A 500-test suite holds 99.2% "
        "coverage, and four real Bedrock-hosted runs show the thesis "
        "directly: the system would rather fail a candidate than invent "
        "experience they lack. The report closes with the challenges this "
        "build surfaced and the work still open.",
    )

    # --- 1. Introduction -------------------------------------------------------
    add_heading(doc, "1. Introduction")
    add_body(
        doc,
        "Resume writing is adversarial-by-default for a language model: the "
        "easiest way to make a draft look stronger is to overstate scope, "
        "seniority, or a metric the evidence doesn't support. TrustResume's "
        "answer is architectural: generation and verification run as "
        "separate agents, so the writer never grades its own work, and every "
        "claim it produces is checked against retrieved evidence by an "
        "independent Trust Harness before any score reaches a user. This "
        "system is a from-scratch reimplementation of an earlier MSAI-699 "
        "capstone (a RAG and multi-agent pipeline on pydantic-ai and Qdrant) "
        "on LangChain, LangGraph, and ChromaDB, a deliberate learning "
        "exercise in the field's dominant agent framework rather than a "
        "redesign; the architecture, the SQLite storage layer, and the "
        "trust/ATS scoring logic carried over unchanged, while how agents are "
        "built and coordinated did not.",
    )
    add_body(
        doc,
        "Earlier weeks of this course layered a measurement discipline on "
        "top of that port the original project never had: an offline "
        "evaluation harness, per-run cost/latency telemetry, role-tiered "
        "models, and a testable per-user identity model. This report "
        "synthesizes that work: methodology (§2), results (§3), challenges "
        "and ethics (§4), and open work (§5).",
    )

    # --- 2. Methodology -------------------------------------------------------
    add_heading(doc, "2. Methodology")

    add_subheading(doc, "2.1 System architecture")
    add_body(
        doc,
        "A run is driven by a single LangGraph StateGraph (Figure 1). Job "
        "Description and Evidence Retrieval (hybrid vector + SQLite-FTS5 "
        "keyword search, user-scoped) each run once; with a cached Candidate "
        "Profile, they feed a quality loop: the Resume Writer drafts, the "
        "Trust Harness classifies every claim as SUPPORTED, "
        "PARTIALLY_SUPPORTED, or UNSUPPORTED against the evidence, and ATS "
        "Evaluation scores keyword coverage against Trust ≥ 90 and ATS ≥ 85. "
        "The loop does not stop the moment a draft passes: deterministic "
        "feedback drives another rewrite regardless, for a configurable "
        "number of rounds (currently one), because a later rewrite can "
        "still improve ATS coverage without spending any more of the Trust "
        "budget. The run then exports whichever draft actually scored best "
        "— a passing draft always beats a failing one, and ties break on "
        "ATS, not Trust, since every passer already cleared that bar.",
    )
    fig1 = make_pipeline_diagram(tmpdir)
    add_figure(doc, fig1, "Figure 1. The generation pipeline: five sequential agents plus a cached sixth, orchestrated by a LangGraph quality loop.")

    add_subheading(doc, "2.2 Key design decisions")
    add_body(
        doc,
        "Storage stayed on SQLite plus ChromaDB rather than Qdrant "
        "(ADR-0001), every search scoped by a user_id filter. Retrieval is "
        "hybrid (ADR-0010): vector and keyword hits fuse by Reciprocal Rank "
        "Fusion, since cosine similarity and BM25 rank aren't comparable. "
        "One catches a paraphrase sharing no vocabulary with its evidence; "
        "the other catches an exact product name the embedder treats as "
        "interchangeable with a competitor. Temperature is pinned to 0 "
        "(ADR-0013), since the headline claim is a deterministic, "
        "code-computed Trust score; models are tiered by role "
        "(extraction/writer/verifier) so a cheap model fills a schema while "
        "a stronger one adjudicates trust. Every route resolves its caller "
        "from a request header (ADR-0014) instead of one hardcoded demo "
        "user, making the isolation boundary testable rather than merely "
        "asserted. An opt-in checkpointing layer (ADR-0015) demonstrates "
        "LangGraph's durable-execution capability: a crashed run resumes "
        "from its last completed node instead of re-paying for "
        "already-completed LLM calls, off by default so the pipeline above "
        "is unchanged unless a caller opts in. The quality loop's stopping "
        "rule changed since the original port (ADR-0016): it no longer ends "
        "the instant a draft passes, since a later rewrite can still find a "
        "better-tailored one without spending any more of the Trust budget "
        "a passing draft already has; the run keeps whichever draft "
        "actually scores best rather than whichever one happened to pass "
        "first.",
    )

    add_subheading(doc, "2.3 Development process")
    add_body(
        doc,
        "Development proceeded as a module-by-module port checked against "
        "the original repository's behavior, with one exception: adopting "
        "LangChain-idiomatic constructs where they were a clear win. An AI "
        "coding assistant (Claude Code) was used throughout: for the port, "
        "the measurement layer, running evaluations, and drafting reports. "
        "Every change was reviewed against the running suite and real "
        "evaluation output before acceptance.",
    )

    add_subheading(doc, "2.4 Testing and evaluation methodology")
    add_body(
        doc,
        "Correctness and quality are measured separately. Pytest runs "
        "entirely offline under a 95% coverage gate; a smaller live-marked "
        "slice exercises the real embedding model and a real HTTP server on "
        "demand. Neither tells you whether the system is getting better at "
        "its job, so a separate offline evaluation harness (ADR-0011) scores "
        "retrieval quality and Trust Harness accuracy against labeled ground "
        "truth. This is deliberately distinct from the product-facing "
        "evaluation package, which scores one resume at runtime and cannot "
        "detect that the grader itself has gone wrong.",
    )

    # --- 3. Results -------------------------------------------------------
    add_heading(doc, "3. Results")

    add_subheading(doc, "3.1 Automated test suite and CI")
    add_body(
        doc,
        "The suite holds 500 passing tests (9 live-marked tests deselected "
        "by default) at 99.2% coverage against a 95% gate. CI runs lint, a "
        "format check, strict mypy, and the full suite on Python 3.11 and "
        "3.13 from a locked dependency file on every push and pull request.",
    )

    add_subheading(doc, "3.2 Retrieval evaluation")
    add_body(
        doc,
        "Scored at k=8, the depth a real generation actually retrieves, "
        "hybrid retrieval reaches recall@8 1.000 and MRR 0.938 over a "
        "10-document labeled corpus and 9 queries, including one query "
        "sharing no vocabulary with its answer and one naming an exact "
        "product the embedder treats as interchangeable with a competitor. "
        "Precision@8 is 0.156, expected rather than a weakness (most queries "
        "have one relevant document, capping precision at 0.125). The one "
        "deliberately unanswerable query still returns a full 8 results, "
        "since retrieval has no notion of “no good answer.” That's a weak "
        "spot discussed in Section 4.",
    )
    add_table(
        doc,
        ["Metric", "Value", "Corpus"],
        [
            ["Recall@8", "1.000", "10 documents / 9 queries"],
            ["MRR", "0.938", "10 documents / 9 queries"],
            ["Precision@8", "0.156", "capped by single-relevant-doc queries"],
            ["Hit rate", "1.000", "N/A"],
        ],
    )

    add_subheading(doc, "3.3 Trust Harness evaluation and calibration fix")
    add_body(
        doc,
        "The Trust Harness's original prompt said only to “be strict... do "
        "not give the draft the benefit of the doubt,” without defining its "
        "three labels. Scored against a 12-case labeled dataset on a real "
        "Bedrock model, that prompt measured accuracy 0.500/macro-F1 0.489, "
        "with every error off by one severity step toward under-crediting. "
        "Even a claim restating its evidence verbatim was downgraded. "
        "Explicit label definitions plus a named asymmetry between error "
        "types (a false pass ships a fabrication; a false flag costs one "
        "rewrite) raised accuracy/macro-F1 to 0.833, reproduced identically "
        "on repeat runs, with dangerous errors held at zero throughout.",
    )

    add_subheading(doc, "3.4 End-to-end generations against a real model")
    add_body(
        doc,
        "One synthetic candidate was run against four job postings, real "
        "Bedrock end to end, each drafted twice — the loop no longer stops "
        "the instant a draft passes (ADR-0016), so even Scenario 1's "
        "already-passing first attempt got a second draft anyway. Three of "
        "four still fail the quality gate. That's deliberate: a gate that "
        "never rejects anything is decoration. Scenario 4 is the clearest "
        "evidence for the project's design bet: given an iOS posting for a "
        "backend candidate, the writer refused to invent iOS experience, so "
        "every claim was supported (Trust 100) and the resume matched no "
        "required keyword (ATS 0), identically on both drafts. The system "
        "chose to fail the candidate rather than lie for them.",
    )
    add_table(
        doc,
        ["Scenario", "Result", "Trust /90", "ATS /85", "Drafts", "LLM calls", "Cost"],
        [
            ["1-strong-match", "PASS", "93", "95", "2", "6", "$0.53"],
            ["2-partial-match", "FAIL", "81", "63", "2", "6", "$0.56"],
            ["3-inflation-pressure", "FAIL", "83", "42", "2", "6", "$0.60"],
            ["4-wrong-domain", "FAIL", "100", "0", "2", "6", "$0.51"],
        ],
    )
    fig2 = make_results_chart(tmpdir)
    add_figure(doc, fig2, "Figure 2. Trust Harness accuracy before/after the calibration fix (left) and Trust/ATS scores across four real Bedrock generations, dashed lines marking the quality gate (right).")
    add_body(
        doc,
        "Scenario 3 shows the checks working under pressure: asked for "
        "staff-level scope the evidence doesn't fully support, the writer "
        "overstated three claims outright — UNSUPPORTED, including "
        "“defining the architectural roadmap for platform-wide event "
        "streaming” — and inflated a fourth to PARTIALLY_SUPPORTED, "
        "rather than staying within what the evidence documents. The Trust "
        "Harness caught all four rather than letting any read as fact. "
        "Because every scenario now runs both of its drafts, cost per run "
        "was uniform this time (six LLM calls, $0.51-$0.60) rather than the "
        "four-to-ten-call spread a lucky-or-unlucky first draft produced "
        "under the old stop-on-first-pass rule — more predictable, though "
        "not free. The second draft's own effect was consistent and not "
        "obviously beneficial: in three of the four scenarios (including "
        "Scenario 1, where both drafts already passed) the rewrite traded "
        "ATS for Trust rather than improving both — Scenario 1 moved from "
        "93/95 to 100/90, still passing but with a worse ATS score, so the "
        "run correctly exported the first draft instead of the second. That "
        "consistent pattern, on a four-scenario sample, is preliminary "
        "evidence rather than proof, but it is why the default rewrite "
        "count was set to one rather than restored to the original three "
        "(ADR-0016) until a larger real-provider comparison says otherwise. "
        "Telemetry (ADR-0012) attributes every one of those calls to its "
        "node and reports an honest null rather than a partial total for "
        "any unpriced model.",
    )

    # --- 4. Discussion -------------------------------------------------------
    add_heading(doc, "4. Discussion")

    add_subheading(doc, "4.1 What worked")
    add_body(
        doc,
        "Treating this as a faithful port kept every regression "
        "attributable to a specific decision. The offline-first test suite "
        "stayed fast enough to run on every change, catching most "
        "structural bugs before a real model call. Most importantly, "
        "separating generation from verification, and both from the "
        "harness grading the verifier itself, is what let §3.3's "
        "calibration bug be found at all. Nothing in the runtime pipeline "
        "could reveal it, since the Trust score is computed from whatever "
        "the harness reports.",
    )

    add_subheading(doc, "4.2 Challenges encountered and how they were addressed")
    add_body(
        doc,
        "Several consequential bugs were invisible to static analysis and "
        "unit tests, surfacing only when the real system ran:",
    )
    add_bullets(
        doc,
        [
            "Trust Harness miscalibration (§3.3): an undefined “be strict” "
            "instruction caused a uniform downgrade. Fixed with explicit "
            "label definitions, verified by an A/B comparison, not "
            "inspection.",
            "Token/cost invisibility: with_structured_output consumes the "
            "AIMessage before usage_metadata reaches the agent. Fixed with a "
            "LangGraph callback that sees every nested model call.",
            "Structured-output brittleness: an empty section heading failed "
            "the strict schema and crashed parsing, losing a paid-for "
            "draft. Fixed with a lenient private schema, relabeled in code.",
            "PDF export crashing on real prose: fpdf2's core font only "
            "encodes Latin-1, and an em dash in real Bedrock output "
            "destroyed an already-paid-for export. Fixed with "
            "transliteration and a degrade-to-‘?’ fallback.",
            "A silent 422 on every route: FastAPI couldn't resolve a type "
            "alias defined inside a function once `from __future__ import "
            "annotations` was active, and only booting the real app caught "
            "it. Fixed by moving it to module scope.",
        ],
    )
    add_body(
        doc,
        "These are integration-boundary bugs, not logic errors inside one "
        "function. That's exactly the class an offline-first suite is "
        "weakest against and a real end-to-end run is strongest against, "
        "the argument for keeping live and manual smoke-test paths at all.",
    )

    add_subheading(doc, "4.3 Ethical considerations")
    add_body(
        doc,
        "Security: externally sourced text, such as a posting, an uploaded "
        "document, or retrieved evidence, is wrapped and marked as data "
        "before reaching a prompt, so injection reads as content, not "
        "instruction. Fairness: the Trust Harness doubles as a fairness "
        "mechanism, since an ungrounded generator would favor candidates "
        "whose backgrounds happen to embellish well; the ATS score carries "
        "a smaller residual bias risk and is treated as candidate feedback, "
        "not a claim about a real employer's system. Privacy: every record "
        "and vector carries a user_id, filtered server-side. The API still "
        "resolves identity from a self-reported header with no "
        "authentication behind it: the clearest ethical gap carried into "
        "Section 5.",
    )

    add_subheading(doc, "4.4 Limitations")
    add_body(
        doc,
        "Two PARTIALLY_SUPPORTED claims still misclassify as UNSUPPORTED. "
        "That's the safe direction, but not a solved problem. Retrieval "
        "always fills "
        "its configured depth even with nothing relevant, relying on the "
        "Trust Harness as the actual downstream safeguard. PDF export "
        "degrades any non-Latin-1 character to ‘?’, and the storage schema "
        "has no migration tool: a change is invisible to a database file "
        "created before it.",
    )

    # --- 5. Future Work -------------------------------------------------------
    add_heading(doc, "5. Future Work")
    add_bullets(
        doc,
        [
            "Real authentication and rate limiting, replacing the "
            "self-reported X-User-Id header. Every store call is already "
            "user-scoped.",
            "A relevance threshold on retrieval, so an unanswerable query "
            "returns fewer chunks instead of always filling top-k.",
            "A larger labeled Trust Harness dataset targeting compound and "
            "seniority-inflation claims, the two remaining failure modes.",
            "Unicode-capable PDF export (a bundled TTF) instead of "
            "degrading non-Latin-1 characters to ‘?’.",
            "Real schema migration tooling, and an incremental scale path "
            "(Kubernetes or a managed service) if demand grows.",
        ],
    )

    # --- 6. Conclusion --------------------------------------------------------------
    add_heading(doc, "6. Conclusion")
    add_body(
        doc,
        "This reimplementation carries forward its predecessor's central "
        "bet: a resume generator is trustworthy only if something other "
        "than the generator checks its claims. The measurement layer built "
        "alongside it is what let that bet be verified rather than "
        "assumed. The evaluation harness caught a real regression in "
        "the one component the system's credibility rests on, and a real "
        "end-to-end run showed the resulting behavior directly: the system "
        "would rather fail an ATS check than fabricate experience. What "
        "remains open is documented, not hidden: authentication, a "
        "retrieval-side relevance floor, and a dataset gap around two claim "
        "types the harness still gets wrong.",
    )

    # --- AI Disclosure Statement -----------------------------------------------------
    add_para(doc, "AI Disclosure Statement", bold=True)
    add_body(
        doc,
        "Throughout this project, an AI coding assistant (Claude Code) was "
        "used to implement the LangChain/LangGraph/ChromaDB port, build and "
        "run the evaluation harness and telemetry layer, diagnose and fix "
        "the bugs in Section 4.2, generate the sample runs in Section 3.4, "
        "and draft this report and its figures. All numeric results came "
        "from actually running the project's code: the test suite, the "
        "evals CLI, or a real Bedrock-hosted model. None were estimated. "
        "All findings and this report's final text were reviewed and confirmed "
        "by the author.",
    )

    return doc


def main() -> None:
    with tempfile.TemporaryDirectory() as tmp:
        doc = build(Path(tmp))
        OUTPUT.parent.mkdir(parents=True, exist_ok=True)
        doc.save(OUTPUT)
    print(f"Wrote {OUTPUT.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
