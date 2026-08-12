"""Generate docs/weekly-submissions/Week6_Model_Testing_and_Debugging_Report.docx.

Throwaway doc-tooling, like build_design_pdf.py/generate_sample_documents.py —
not part of the app, not covered by the test suite. Mirrors the formatting of
the (hand-authored, in Word) Week 5 report so the two look like one series:
Times New Roman, double-spaced, centered title page, bold numbered section
headings, "List Bullet" for bullet lists.

    .venv/bin/python scripts/build_week6_report.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "weekly-submissions" / "Week6_Model_Testing_and_Debugging_Report.docx"

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)


def _set_run_font(run) -> None:  # type: ignore[no-untyped-def]
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE


def add_para(doc: Document, text: str = "", *, center: bool = False, bold: bool = False) -> None:  # type: ignore[no-untyped-def]
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        run = p.add_run(text)
        run.bold = bold
        _set_run_font(run)


def add_body(doc: Document, text: str) -> None:  # type: ignore[no-untyped-def]
    """A double-spaced body paragraph, left-aligned (Word's default justify-none)."""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 2.0
    run = p.add_run(text)
    _set_run_font(run)


def add_heading(doc: Document, text: str) -> None:  # type: ignore[no-untyped-def]
    add_para(doc, text, bold=True)


def add_bullets(doc: Document, items: list[str]) -> None:  # type: ignore[no-untyped-def]
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 2.0
        run = p.add_run(item)
        _set_run_font(run)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:  # type: ignore[no-untyped-def]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for cell, text in zip(table.rows[0].cells, headers, strict=True):
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        _set_run_font(run)
    for row, values in zip(table.rows[1:], rows, strict=True):
        for cell, text in zip(row.cells, values, strict=True):
            run = cell.paragraphs[0].add_run(text)
            _set_run_font(run)
    doc.add_paragraph()


def build() -> Document:  # type: ignore[no-untyped-def]
    doc = Document()

    for style_name in ("Normal", "List Bullet"):
        style = doc.styles[style_name]
        style.font.name = FONT_NAME
        style.font.size = FONT_SIZE

    # --- Title page ---------------------------------------------------------
    add_para(doc, "", center=True)
    add_para(doc, "MSAI 699 Capstone Project", center=True)
    add_para(doc, "", center=True)
    add_para(doc, "Instructor: Dr. Gamini Bulumulle", center=True)
    add_para(doc, "", center=True)
    add_para(doc, "Week 6: Testing, Evaluation & Documentation: Model Testing and Debugging Report", center=True)
    add_para(doc, "", center=True)
    add_para(doc, "Peng Fei Xu", center=True)
    add_para(doc, "", center=True)
    add_para(doc, "University of the Cumberlands", center=True)
    add_para(doc, "", center=True)
    add_para(doc, "Submitted to the University of the Cumberlands", center=True)
    add_para(doc, "in Partial Fulfillment of the Requirements of the Degree of", center=True)
    add_para(doc, "Master of Science in Artificial Intelligence", center=True)
    add_para(doc, "", center=True)
    add_para(doc, "Aug 7, 2026", center=True)
    doc.add_page_break()

    # --- Abstract ------------------------------------------------------------
    add_para(doc, "Abstract", center=True, bold=True)
    add_body(
        doc,
        "This report documents an A/B test, error analysis, and reliability pass on "
        "TrustResume's Trust Harness, the agent that checks every resume claim against "
        "the candidate's own evidence. Scored against a 12-case labeled dataset, the "
        "harness's original prompt measured accuracy 0.500 / macro-F1 0.489, with every "
        "error running the same direction — an undefined “be strict” instruction "
        "causing a uniform one-notch downgrade. Replacing it with explicit label "
        "definitions and a named error asymmetry raised accuracy/macro-F1 to 0.833, "
        "reproduced identically across independent runs at temperature 0, with zero "
        "dangerous (too-lenient) errors before and after. Two remaining "
        "misclassifications and a retrieval weak spot are analyzed as open findings. "
        "Separately, this pass added per-step cost telemetry and fixed a "
        "pricing-configuration gap that was silently reporting every Bedrock "
        "generation's cost as unknown.",
    )

    # --- 1. Introduction -------------------------------------------------------
    add_heading(doc, "1. Introduction")
    add_body(
        doc,
        "TrustResume's central claim is that every statement in a generated resume is "
        "checked against the candidate's own evidence. The runtime Trust score is "
        "computed directly from what the Trust Harness agent reports, so a harness "
        "that rubber-stamps everything SUPPORTED would produce a perfect score and an "
        "invisible failure. ADR-0011, “An Offline Evaluation Harness with "
        "Labeled Ground Truth,” added that check: a hand-labeled claim/evidence "
        "dataset scored offline, independent of any live user.",
    )

    # --- 2. Testing Methodology -------------------------------------------------
    add_heading(doc, "2. Testing Methodology")
    add_body(
        doc,
        "TrustHarnessAgent extracts every factual claim from a draft and classifies "
        "each as SUPPORTED, PARTIALLY_SUPPORTED, or UNSUPPORTED. The 12-case dataset "
        "(evals/datasets/trust_claims.jsonl) is weighted toward that boundary, where a "
        "resume actually lies. python -m trustresume.evals --suite trust drives every "
        "case through the real agent and model (Bedrock, claude-opus-4-6, temperature "
        "0) and scores accuracy, macro-F1, and a confusion matrix; macro-F1 guards "
        "against a harness that never predicts UNSUPPORTED looking accurate while "
        "failing its only job. A dangerous_errors count is tracked separately, since a "
        "too-lenient error ships a fabrication while a too-strict one only costs one "
        "rewrite. The A/B test below holds model, temperature, and dataset fixed and "
        "varies only the prompt; reproducibility was confirmed by running the current "
        "configuration twice independently, returning identical results both times.",
    )

    # --- 3. A/B Test Results -----------------------------------------------------
    add_heading(doc, "3. A/B Test Results: Trust Harness Prompt")
    add_body(
        doc,
        "The original prompt said only “be strict... do not give the draft the "
        "benefit of the doubt,” without ever defining the three labels:",
    )
    add_table(
        doc,
        ["Metric", "Before (original)", "After (redefined labels)"],
        [
            ["Accuracy", "0.500", "0.833"],
            ["Macro-F1", "0.489", "0.833"],
            ["SUPPORTED recall", "0.25", "1.00"],
            ["Dangerous (too-lenient) errors", "0", "0"],
        ],
    )
    add_body(
        doc,
        "All six errors under the original prompt ran the same direction, off by "
        "exactly one severity step toward under-crediting — e.g. a claim restating "
        "its evidence verbatim was downgraded to PARTIALLY_SUPPORTED. That is a model "
        "told to be cautious with no defined boundary, not one that cannot "
        "distinguish supported from unsupported. The fix replaced the instruction "
        "with explicit definitions (SUPPORTED covers restatements and entailed "
        "generalizations; PARTIALLY_SUPPORTED is reserved for overstated scope, "
        "seniority, or numbers; UNSUPPORTED for claims with no basis at all) and named "
        "the asymmetry directly, so the model resists overstatement without "
        "discounting everything. Both conditions were reproduced independently this "
        "session — 0.833/0.833 twice, and reverting to the original prompt returned "
        "0.500/0.489, exactly matching. Dangerous errors held at zero throughout: the "
        "fix gained accuracy without becoming more permissive.",
    )

    # --- 4. Error Analysis --------------------------------------------------------
    add_heading(doc, "4. Error Analysis")
    add_body(
        doc,
        "Two cases still misclassify after the fix, both PARTIALLY_SUPPORTED claims "
        "judged UNSUPPORTED — the safe direction, but worth naming precisely:",
    )
    add_table(
        doc,
        ["Case", "Claim", "Evidence", "Expected → Predicted"],
        [
            [
                "t6",
                "Owned CI/CD and cut build times by 90%.",
                "Owned CI/CD for 40 engineers; a caching layer cut build time "
                "11→4 min (~64%).",
                "PARTIAL → UNSUPPORTED",
            ],
            [
                "t7",
                "Expert-level Kubernetes administrator.",
                "Self-taught enough to debug one scheduling bug.",
                "PARTIAL → UNSUPPORTED",
            ],
        ],
    )
    add_body(
        doc,
        "t6 is compound (true ownership, false 90% figure); the model let the false "
        "part drag the whole claim to UNSUPPORTED instead of averaging. t7 is "
        "seniority inflation the prompt's own definitions almost describe verbatim as "
        "PARTIALLY_SUPPORTED, yet the model called it baseless — a genuinely "
        "borderline case still resolved toward the stricter neighbor. A concrete next "
        "step is adding more compound-claim and seniority-gap examples to the labeled "
        "dataset itself. A related weak spot: the retrieval eval's one deliberately "
        "unanswerable query still returns a full 8 irrelevant chunks, since hybrid "
        "retrieval has no “no good answer” notion. The Trust Harness is the "
        "actual safeguard here (Section 3's zero dangerous errors is evidence it "
        "works), but a retrieval-side relevance threshold would be a cheaper first "
        "line of defense.",
    )

    # --- 5. Reliability Improvements & Best Practices -----------------------------
    add_heading(doc, "5. Reliability Improvements & Best Practices")
    add_body(
        doc,
        "Two changes made alongside the fix address reliability, both found by "
        "measuring a real run rather than trusting it worked. Telemetry previously "
        "recorded tokens-per-model and duration-per-node as two unlinked lists; "
        "LangGraph already tags each LLM call with the node that made it, so the "
        "tracker now attributes tokens and cost to that node directly, with no agent "
        "changes required, and every run now emits a metrics.json with a per-step "
        "breakdown. Separately, config/pricing.json listed OpenAI rates only, so every "
        "Bedrock generation reported cost_usd = null — correct behavior for an "
        "unpriced model, but Bedrock was never actually unpriced, just missing. The "
        "fix adds real Claude rates plus a version-extension guard so a future model "
        "release cannot silently inherit an earlier one's price.",
    )
    add_body(
        doc,
        "Best practices: never report a partial/default number when the true value is "
        "unknown; when a classifier's errors are correlated, examine their direction "
        "before touching the model or data; treat too-lenient and too-strict errors "
        "as different severities; reproduce a before/after comparison at a pinned "
        "temperature, in both directions, before treating a fix as confirmed.",
    )

    # --- 6. Conclusion --------------------------------------------------------------
    add_heading(doc, "6. Conclusion")
    add_body(
        doc,
        "The offline evaluation measured a real regression (accuracy 0.500), "
        "root-caused it to an undefined strictness instruction, and confirmed a fix "
        "(accuracy 0.833) that closed the gap without opening the more dangerous "
        "false-pass failure mode. Two misclassifications and a retrieval weak spot "
        "remain as documented open findings rather than treated as solved. Alongside "
        "the model-quality work, a telemetry and pricing fix closed a gap where every "
        "real generation's cost was silently unreportable, extending the same "
        "“never report an unearned number” discipline to cost accounting.",
    )

    # --- AI Disclosure Statement -----------------------------------------------------
    add_para(doc, "AI Disclosure Statement", bold=True)
    add_body(
        doc,
        "In August 2026, during Week 6 of this project, an AI coding assistant "
        "(Claude) was used to run the evaluation harness, independently reproduce the "
        "A/B comparison, implement the Section 5 fix, and draft this report. The "
        "Trust Harness prompt fix (Section 3) was implemented in an earlier session "
        "and is reported here, not newly authored. All numeric results came from "
        "running the project's code against a real Bedrock-hosted model. All findings "
        "were reviewed and confirmed by the author.",
    )

    return doc


def main() -> None:
    doc = build()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
