"""Generate docs/weekly-submissions/Week8_Code_Repository_Report.docx.

Throwaway doc-tooling, like build_week7_report.py. Not part of the app, not
covered by the test suite. Documents the source code submitted for the
capstone: what is included, where each part came from, whether it runs, and
what challenges surfaced during development. Written in the first person
throughout: this rubric assumes a group project, but this capstone was
completed individually, so "the group" is rewritten as "I" rather than kept
verbatim.

    .venv/bin/python scripts/build_week8_code_repository.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "weekly-submissions" / "Week8_Code_Repository_Report.docx"

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)


def _set_run_font(run) -> None:  # type: ignore[no-untyped-def]
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE


def _no_extra_spacing(p) -> None:  # type: ignore[no-untyped-def]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 2.0


def add_para(doc: Document, text: str = "", *, center: bool = False, bold: bool = False) -> None:  # type: ignore[no-untyped-def]
    p = doc.add_paragraph()
    _no_extra_spacing(p)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        run = p.add_run(text)
        run.bold = bold
        _set_run_font(run)


def add_body(doc: Document, text: str) -> None:  # type: ignore[no-untyped-def]
    p = doc.add_paragraph()
    _no_extra_spacing(p)
    run = p.add_run(text)
    _set_run_font(run)


def add_heading(doc: Document, text: str) -> None:  # type: ignore[no-untyped-def]
    add_para(doc, text, bold=True)


def add_bullets(doc: Document, items: list[str]) -> None:  # type: ignore[no-untyped-def]
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        _no_extra_spacing(p)
        run = p.add_run(item)
        _set_run_font(run)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:  # type: ignore[no-untyped-def]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for cell, text in zip(table.rows[0].cells, headers, strict=True):
        _no_extra_spacing(cell.paragraphs[0])
        cell.paragraphs[0].paragraph_format.line_spacing = 1.0
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        _set_run_font(run)
    for row, values in zip(table.rows[1:], rows, strict=True):
        for cell, text in zip(row.cells, values, strict=True):
            _no_extra_spacing(cell.paragraphs[0])
            cell.paragraphs[0].paragraph_format.line_spacing = 1.0
            run = cell.paragraphs[0].add_run(text)
            _set_run_font(run)
    p = doc.add_paragraph()
    _no_extra_spacing(p)


def build() -> Document:  # type: ignore[no-untyped-def]
    doc = Document()

    for style_name in ("Normal", "List Bullet"):
        style = doc.styles[style_name]
        style.font.name = FONT_NAME
        style.font.size = FONT_SIZE

    section = doc.sections[0]
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # --- Title page ---------------------------------------------------------
    add_para(doc, "", center=True)
    add_para(doc, "MSAI 699 Capstone Project", center=True)
    add_para(doc, "", center=True)
    add_para(doc, "Instructor: Dr. Gamini Bulumulle", center=True)
    add_para(doc, "", center=True)
    add_para(doc, "Week 8: Code Repository", center=True)
    add_para(doc, "", center=True)
    add_para(doc, "Peng Fei Xu", center=True)
    add_para(doc, "", center=True)
    add_para(doc, "University of the Cumberlands", center=True)
    add_para(doc, "", center=True)
    add_para(doc, "Submitted to the University of the Cumberlands", center=True)
    add_para(doc, "in Partial Fulfillment of the Requirements of the Degree of", center=True)
    add_para(doc, "Master of Science in Artificial Intelligence", center=True)
    add_para(doc, "", center=True)
    add_para(doc, "Aug 11, 2026", center=True)
    doc.add_page_break()

    # --- Intro ------------------------------------------------------------
    add_para(doc, "Code Repository Report", center=True, bold=True)
    add_body(
        doc,
        "This report documents the source code submitted for the "
        "TrustResume capstone: what is included, where each part came "
        "from, whether it runs, and what challenges came up while "
        "building it. The assignment rubric is written for a group "
        "project, but I completed this capstone individually, so every "
        "file in the repository was written or ported by me alone; the "
        "sections below use \"I\" throughout rather than \"the group.\"",
    )

    # --- 1. Repository -------------------------------------------------------
    add_heading(doc, "1. Repository")
    add_body(
        doc,
        "The full source code lives at "
        "https://github.com/pxu/trustresume_langchain, a repository under "
        "my personal GitHub account. It is a Python project managed with "
        "uv, with a committed lock file "
        "(uv.lock) so an installation from it reproduces the exact "
        "dependency versions used during development. A GitHub Actions "
        "workflow runs lint, a format check, strict type checking, and the "
        "full test suite on every push, on Python 3.11 and 3.13.",
    )

    # --- 2. What's included -------------------------------------------------
    add_heading(doc, "2. What Is Included")
    add_body(
        doc,
        "The repository is organized by responsibility. The table below "
        "summarizes the main directories submitted with this report.",
    )
    add_table(
        doc,
        ["Directory", "Contents"],
        [
            ["src/trustresume/models", "Shared data schemas used by every other package"],
            ["src/trustresume/storage", "SQLite repositories for users, documents, jobs, chunks, and resumes"],
            ["src/trustresume/retrieval", "Embeddings, the Chroma vector store, and hybrid vector plus keyword search"],
            ["src/trustresume/ingestion", "Document parsing, cleaning, chunking, and duplicate detection"],
            ["src/trustresume/agents", "The six agents that extract, retrieve, write, verify, and score a resume"],
            ["src/trustresume/orchestration", "The LangGraph orchestrator that runs the quality loop"],
            ["src/trustresume/trust_verification", "Prompt and scoring logic for the Trust Harness"],
            ["src/trustresume/evaluation", "ATS keyword coverage scoring shown to the user"],
            ["src/trustresume/evals", "The offline evaluation harness scored against labeled datasets"],
            ["src/trustresume/telemetry.py", "Per run token, cost, and latency tracking"],
            ["src/trustresume/export", "Markdown and PDF rendering of a generated resume"],
            ["src/trustresume/api", "The FastAPI backend and the provider agnostic model factory"],
            ["src/trustresume/ui", "The Streamlit frontend"],
            ["src/trustresume/poc", "A manual, credential requiring smoke test script, not part of the app"],
            ["tests/", "474 automated unit and integration tests"],
            ["evals/", "Labeled datasets and the recorded evaluation baseline"],
            ["docs/architecture", "Design documents and architecture decision records"],
            ["scripts/", "One off tooling, including the scripts that generated this report series"],
            ["config/", "LLM provider selection and pricing configuration"],
        ],
    )

    # --- 3. Attribution -------------------------------------------------------
    add_heading(doc, "3. Source Code Attribution")
    add_body(
        doc,
        "This project is a rebuild of an earlier capstone of mine, "
        "trustresume (github.com/pxu/trustresume), which used pydantic-ai "
        "and Qdrant instead of LangChain, LangGraph, and ChromaDB. The "
        "storage layer in that original project was carried over nearly "
        "unchanged into this one, and the overall shape of the system, "
        "six agents behind an orchestrated quality loop with trust "
        "verification kept separate from generation, was reapplied "
        "conceptually onto the new stack rather than copied line for "
        "line. docs/architecture/high-level-design.md and the "
        "architecture decision records in that same folder document, "
        "module by module, what was ported as is and what was rebuilt.",
    )
    add_body(
        doc,
        "Beyond that, the project depends on standard open source "
        "libraries declared in pyproject.toml: LangChain and LangGraph for "
        "agent orchestration, ChromaDB for vector storage, FastAPI and "
        "Streamlit for the backend and frontend, unstructured for document "
        "parsing, fastembed for embeddings, fpdf2 for PDF export, and "
        "boto3 plus the LangChain provider packages for AWS Bedrock, "
        "OpenAI, and Google. These are installed dependencies, not code "
        "copied into the repository. Agent orchestration in this codebase "
        "follows patterns shown in LangChain's and LangGraph's own "
        "documentation, structured output binding and the StateGraph "
        "abstraction in particular, rather than any specific blog post or "
        "third party tutorial. No external tutorial code was copied into "
        "this repository.",
    )
    add_body(
        doc,
        "An AI coding assistant, Claude Code, was used throughout "
        "implementation, testing, and documentation. Its role is spelled "
        "out further in the AI Disclosure Statement at the end of this "
        "report.",
    )

    # --- 4. Execution status -------------------------------------------------
    add_heading(doc, "4. Execution Status")
    add_body(
        doc,
        "The automated test suite runs with no credentials and no network "
        "access, and it passes: 474 tests, 99.4% statement and branch "
        "coverage against a 95% gate, verified with pytest immediately "
        "before this report was written. Everything that needs a real "
        "external service is isolated behind a flag or a separate marker, "
        "so grading the bulk of this submission does not require API "
        "keys of any kind.",
    )
    add_table(
        doc,
        ["Component", "Runs without credentials", "Notes"],
        [
            ["pytest (default)", "Yes", "474 tests, 99.4% coverage, in memory database and vector store"],
            ["pytest -m live", "Partly", "needs network access for an embedding model download and a real local server, but no API key"],
            ["API, test provider", "Yes", "TRUSTRESUME_LLM_PROVIDER=test drives the full pipeline with a synthetic model"],
            ["API, Bedrock, OpenAI, or Google", "No", "needs real provider credentials, not included in this submission"],
            ["scripts/manual_rag_test.py", "No", "a manual smoke test against a real Bedrock model"],
            ["src/trustresume/poc/llm_smoke_test.py", "No", "a manual, credential requiring script, excluded from the coverage gate by design"],
            ["Docker Compose", "Yes", "defaults to the offline test provider"],
        ],
    )
    add_body(
        doc,
        "In short, nothing in this submission fails to execute for a "
        "reason I could not explain. The parts that will not run without "
        "extra setup are the ones that call a real, paid LLM provider, "
        "and that is expected: this project was built offline first on "
        "purpose, exactly so a grader without my AWS or OpenAI credentials "
        "can still run and verify the great majority of it.",
    )

    # --- 5. Challenges -------------------------------------------------------
    add_heading(doc, "5. Challenges Encountered and Their Causes")
    add_body(
        doc,
        "Several of the most consequential problems in this project were "
        "invisible to my test suite and only surfaced when I ran the real "
        "system end to end. Each one below is a symptom, then what I "
        "believe actually caused it.",
    )
    add_bullets(
        doc,
        [
            "The Trust Harness, the agent that checks every resume claim "
            "against evidence, scored only 50 percent accuracy on a "
            "labeled test set. The cause was a prompt that said to be "
            "strict without ever defining what strict meant, so the model "
            "applied a uniform one notch downgrade to nearly every claim "
            "rather than judging each one on its own merits.",
            "Token counts and cost were invisible for every generation. "
            "The cause is that LangChain's structured output binding "
            "returns a parsed object and consumes the underlying model "
            "response, including its usage data, before that response "
            "ever reaches the calling code. I fixed it with a callback "
            "attached once per run that intercepts every model call "
            "before it is consumed.",
            "PDF export crashed and destroyed an entire generation, after "
            "every model call in that run had already been paid for. The "
            "cause is that the PDF library's default font only supports "
            "Latin-1 text, and a real model's output routinely contains "
            "an em dash or a curly quote, characters my earlier, "
            "synthetic test fixtures never happened to include.",
            "Every API route started returning an unrecognized parameter "
            "error with no obvious cause. The cause turned out to be a "
            "Python feature, deferred evaluation of type hints, combined "
            "with a type alias I had defined inside a function body "
            "instead of at module level, which made that alias invisible "
            "to FastAPI's own resolution of the annotation. Neither my "
            "type checker nor my unit tests caught this, since they never "
            "actually booted the live application the way a real request "
            "does.",
            "Keyword search crashed whenever a real job posting contained "
            "ordinary punctuation such as a hyphen, a colon, or a slash. "
            "The cause is that SQLite's full text search treats those "
            "characters as query operators, not literal text, and a real "
            "job posting reliably contains at least one of them.",
            "Parsing a single resume file was taking close to a minute. "
            "The cause is that the parsing library's default strategy "
            "loads a large layout detection model meant for scanned or "
            "image heavy documents, which is unnecessary overhead for "
            "plain text extraction from a resume.",
            "Every real generation against my default provider quietly "
            "reported its cost as unknown. The cause is that my pricing "
            "configuration file was written and tested against one "
            "provider's models and was never extended to include the "
            "model I actually use by default, so the code's fallback "
            "behavior, report nothing rather than guess, masked the gap "
            "instead of surfacing it.",
        ],
    )
    add_body(
        doc,
        "All seven share a pattern. None of them were logic errors inside "
        "a single function that a unit test would naturally catch. Each "
        "one lived at the boundary between a framework's abstraction and "
        "an assumption I had made about it, and each one only showed up "
        "once I ran the real system against real data or a real model. "
        "That is the reason I kept a slower, credential requiring testing "
        "path in this project instead of trusting my fast, offline test "
        "suite alone.",
    )

    # --- 6. Conclusion -------------------------------------------------------
    add_heading(doc, "6. Conclusion")
    add_body(
        doc,
        "The submitted repository contains every line of source code "
        "behind this capstone, written or ported by me individually, with "
        "its origins credited above. The automated test suite runs "
        "without credentials and passes in full. The handful of scripts "
        "that need a real LLM provider are clearly marked and will not "
        "run without credentials I have not included, which is expected "
        "given how this project was designed. The challenges section "
        "above is meant to save a grader, or a future version of me, the "
        "time of rediscovering the same seven root causes from scratch.",
    )

    # --- AI Disclosure Statement -----------------------------------------------------
    add_para(doc, "AI Disclosure Statement", bold=True)
    add_body(
        doc,
        "An AI coding assistant, Claude Code, was used throughout this "
        "project to help implement the source code described above, "
        "write and run its test suite, diagnose the causes described in "
        "Section 5, and draft this report. Every change was reviewed "
        "against the running test suite before I accepted it, and the "
        "descriptions of what runs and what does not in Section 4 came "
        "from actually running the commands listed there, not from "
        "assumption.",
    )

    return doc


def main() -> None:
    doc = build()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
