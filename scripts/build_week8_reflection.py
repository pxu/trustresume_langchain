"""Generate docs/weekly-submissions/Week8_Personal_Reflection.docx.

Throwaway doc-tooling, like build_week7_report.py. Not part of the app, not
covered by the test suite. A first-person reflection essay (~500 words), not
a technical report: no headings, no tables, no figures, just an
introduction, body, and conclusion, in the same title-page format as the
rest of the weekly-submissions series.

    .venv/bin/python scripts/build_week8_reflection.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "weekly-submissions" / "Week8_Personal_Reflection.docx"

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(12)


def _set_run_font(run) -> None:  # type: ignore[no-untyped-def]
    run.font.name = FONT_NAME
    run.font.size = FONT_SIZE


def _no_extra_spacing(p) -> None:  # type: ignore[no-untyped-def]
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 2.0


def add_para(doc: Document, text: str = "", *, center: bool = False, bold: bool = False) -> None:  # type: ignore[no-untyped-def]
    p = doc.add_paragraph()
    _no_extra_spacing(p)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if text:
        run = p.add_run(text)
        run.bold = bold
        _set_run_font(run)


def add_body(doc: Document, text: str) -> None:  # type: ignore[no-untyped-def]
    p = doc.add_paragraph()
    _no_extra_spacing(p)
    run = p.add_run(text)
    _set_run_font(run)


def build() -> Document:  # type: ignore[no-untyped-def]
    doc = Document()

    for style_name in ("Normal", "List Bullet"):
        style = doc.styles[style_name]
        style.font.name = FONT_NAME
        style.font.size = FONT_SIZE

    section = doc.sections[0]
    section.left_margin = Pt(72)
    section.right_margin = Pt(72)
    section.top_margin = Pt(72)
    section.bottom_margin = Pt(72)

    # --- Title page ---------------------------------------------------------
    add_para(doc, "", center=True)
    add_para(doc, "MSAI 699 Capstone Project", center=True)
    add_para(doc, "", center=True)
    add_para(doc, "Instructor: Dr. Gamini Bulumulle", center=True)
    add_para(doc, "", center=True)
    add_para(doc, "Week 8: Personal Reflection", center=True)
    add_para(doc, "", center=True)
    add_para(doc, "Peng Fei Xu", center=True)
    add_para(doc, "", center=True)
    add_para(doc, "University of the Cumberlands", center=True)
    add_para(doc, "", center=True)
    add_para(doc, "Submitted to the University of the Cumberlands", center=True)
    add_para(doc, "in Partial Fulfillment of the Requirements of the Degree of", center=True)
    add_para(doc, "Master of Science in Artificial Intelligence", center=True)
    add_para(doc, "", center=True)
    add_para(doc, "Aug 11, 2026", center=True)
    doc.add_page_break()

    add_para(doc, "Personal Reflection", center=True, bold=True)

    # --- Introduction ---------------------------------------------------------
    add_body(
        doc,
        "I came into this capstone with no real hands-on experience "
        "building AI agents, and I chose the project specifically to close "
        "that gap: to get close enough to industry best practices to walk "
        "into an AI engineer interview speaking from experience, not just "
        "theory. I started with pydantic-ai, Qdrant, and a hand-rolled RAG "
        "pipeline, but kept writing by hand pieces that I later found "
        "LangChain and LangGraph already ship as tested, standard building "
        "blocks (text chunking for embedding was the clearest case). That "
        "gap between what I was reinventing and what the ecosystem already "
        "solved well is what pushed me to rebuild the system on LangChain, "
        "LangGraph, and ChromaDB instead. What I hadn't expected was that "
        "the harder lesson waiting on the other side of that switch had "
        "nothing to do with the framework at all: a system built to check "
        "whether an AI is telling the truth has to be checked itself, just "
        "as rigorously, or its central promise is just another unverified "
        "claim.",
    )

    # --- Body ---------------------------------------------------------
    add_body(
        doc,
        "Technically, I learned a great deal about how LangChain and "
        "LangGraph structure an agentic pipeline: structured-output "
        "binding instead of a hand-rolled agent loop, a StateGraph instead "
        "of a while loop, callbacks as the only place certain information "
        "(like token usage) is still visible once a chain has consumed it. "
        "But the more lasting lesson was conceptual: correctness and "
        "quality are different questions, and a passing test suite answers "
        "neither about a model's actual behavior. My offline evaluation "
        "harness, built to score the system against labeled ground truth, "
        "independent of any live user, caught something my 500 unit tests "
        "never could: the Trust Harness, the one agent this project's "
        "credibility rests on, was miscalibrated. Its prompt said to “be "
        "strict” without ever defining what that meant, so it downgraded "
        "almost every claim by one notch, true or not. Finding that with "
        "data, not intuition, and fixing it with explicit definitions "
        "rather than a vaguer instruction, was the moment this project "
        "stopped feeling like an exercise and started feeling like "
        "engineering.",
    )
    add_body(
        doc,
        "The other challenges I ran into shared a pattern I didn't expect: "
        "the most consequential bugs were invisible to static analysis and "
        "to my offline test suite, only appearing when I ran the real "
        "system against a real model. A resume's em dash crashed PDF export "
        "after every agent call in that run had already been paid for. A "
        "type alias defined in the wrong scope silently broke every API "
        "route, caught only by booting the app, not by mypy. What helped "
        "was asking why my fast, offline-first tests hadn't caught it, and "
        "keeping a slower, real-dependency path specifically for that "
        "blind spot instead of trusting speed alone. I also leaned heavily "
        "on an AI coding assistant throughout this build, which taught me a "
        "different skill than writing code myself: reviewing someone "
        "else's work, or something else's, critically enough to catch "
        "what it missed, the same discipline the Trust Harness applies to "
        "a resume.",
    )

    # --- Conclusion ---------------------------------------------------------
    add_body(
        doc,
        "I started this project wanting to learn a new framework. I finished "
        "it with a much more durable habit: whenever I build something that "
        "makes a claim about quality, correctness, or trust, I now ask what "
        "would catch that claim if it were wrong, and I build that check "
        "before I believe the result. That question feels bigger than this "
        "one capstone, and it's the part I expect to carry into whatever I "
        "build next.",
    )

    return doc


def main() -> None:
    doc = build()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote {OUTPUT.relative_to(ROOT)}")


if __name__ == "__main__":
    main()
