"""Generate the synthetic sample résumés in ``data/sample_documents/``.

These files exist so that `.docx` and `.pdf` parsing is exercised against
*real* binary formats — `unstructured` behaves differently on a genuine Word
file than on anything a test can fabricate inline (see the ``live``-marked
tests in ``tests/unit/test_ingestion.py``) — and so
``scripts/manual_rag_test.py`` has believable evidence to retrieve against.

They replaced two real personal résumés that previously sat here. Real
candidate documents must never live in this repo, even gitignored: an ignore
rule is one `git add -f` or one tooling change away from failing, and the
project's own design scopes candidate data as local-only and user-isolated
(ADR-0001). Generating them from a script instead makes the content
reviewable in a diff and regenerable, rather than an opaque binary nobody can
audit.

Everything below is invented. Names are fictional, emails use the
IANA-reserved ``example.com`` domain (RFC 2606), and phone numbers use the
``555-01xx`` range reserved for fiction — so nothing here can route to a real
person even by accident.

    python scripts/generate_sample_documents.py

Content is written to be a plausible-but-imperfect match for
``data/sample_job_descriptions/`` — deliberately imperfect, because a résumé
that already matches a posting perfectly gives the quality loop nothing to do
and makes a demo look better than the system is.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from fpdf import FPDF
from fpdf.enums import XPos, YPos

OUTPUT_DIR = Path(__file__).resolve().parents[1] / "data" / "sample_documents"

# --- Persona 1: AI engineer (.docx) ---------------------------------------

AI_ENGINEER: list[tuple[str, list[str]]] = [
    (
        "",
        [
            "Jordan Rivera",
            "AI / Machine Learning Engineer",
            "jordan.rivera@example.com | (555) 0142 | Portland, OR",
        ],
    ),
    (
        "Summary",
        [
            "Machine learning engineer with 6 years building and operating "
            "production ML and NLP systems. Comfortable owning a model from "
            "data collection through serving, monitoring, and the eventual "
            "decision to retire it.",
        ],
    ),
    (
        "Experience",
        [
            "Senior ML Engineer, Larkfield Analytics (2022 - present)",
            "Built a retrieval-augmented question answering service over "
            "180,000 internal support tickets. Chunked and embedded the corpus "
            "with sentence-transformers, served answers from a FastAPI backend "
            "over a managed vector index, and added a citation panel so agents "
            "could check the source before trusting an answer.",
            "Cut median answer latency from 4.1s to 900ms by caching embeddings "
            "for repeated queries and moving reranking behind a feature flag "
            "that we could disable during incidents.",
            "Designed the offline evaluation set (600 labeled question/answer "
            "pairs) that gated every prompt change. Two proposed prompt "
            "rewrites were rejected because recall dropped, which would not "
            "have been visible from spot checks.",
            "Owned the model serving stack on Kubernetes: autoscaling GPU node "
            "pools, canary deploys, and a rollback path that took under five "
            "minutes end to end.",
            "Mentored two junior engineers through their first production model "
            "launches, including writing the review checklist the team still "
            "uses for model changes.",
            "",
            "Machine Learning Engineer, Cobalt Retail Group (2020 - 2022)",
            "Built a demand forecasting pipeline in Python and Apache Airflow "
            "covering 12,000 SKUs across 40 stores. Reduced weekly forecast "
            "error (WAPE) from 31% to 19% by adding promotion calendars and "
            "local weather as features.",
            "Replaced a nightly batch scoring job with an incremental pipeline, "
            "cutting compute cost roughly in half and making same-day "
            "replenishment decisions possible for the first time.",
            "Instrumented training and inference with MLflow so that any "
            "production prediction could be traced back to the exact dataset, "
            "code commit, and hyperparameters that produced it.",
            "",
            "Data Scientist, Cobalt Retail Group (2019 - 2020)",
            "Analyzed customer churn and built the first internal dashboard the "
            "merchandising team used weekly. Wrote most of it in SQL and pandas; "
            "the modeling was simple and the data cleaning was not.",
        ],
    ),
    (
        "Selected projects",
        [
            "Document classification service: fine-tuned a small transformer to "
            "route inbound documents into 14 categories at 94% accuracy, "
            "replacing a rules engine that had grown to 800 hand-written "
            "conditions.",
            "Internal embedding benchmark: compared four open embedding models "
            "on our own labeled retrieval set rather than public leaderboards. "
            "The smallest model won on our data, which saved noticeable "
            "inference cost.",
        ],
    ),
    (
        "Skills",
        [
            "Languages: Python (primary), SQL, some Go",
            "ML / NLP: PyTorch, scikit-learn, Hugging Face Transformers, "
            "sentence-transformers, LangChain",
            "Data / infra: Airflow, Spark, PostgreSQL, Redis, Docker, "
            "Kubernetes, AWS (S3, Lambda, SageMaker)",
            "Practices: offline evaluation, A/B testing, model monitoring, "
            "MLflow experiment tracking",
        ],
    ),
    (
        "Education",
        [
            "M.S. Computer Science, Cascade State University (2019)",
            "B.S. Statistics, Cascade State University (2017)",
        ],
    ),
]

# --- Persona 2: senior backend engineer (.pdf) -----------------------------

SENIOR_SDE: list[tuple[str, list[str]]] = [
    (
        "",
        [
            "Taylor Nguyen",
            "Senior Software Engineer, Backend & Distributed Systems",
            "taylor.nguyen@example.com | (555) 0187 | Austin, TX",
        ],
    ),
    (
        "Summary",
        [
            "Backend engineer with 9 years on high-throughput distributed "
            "systems, mostly in logistics and payments. Happiest owning a "
            "service end to end: design, deploy, on-call, and the postmortem "
            "when it breaks.",
        ],
    ),
    (
        "Experience",
        [
            "Senior Software Engineer, Meridian Fulfillment (2021 - present)",
            "Designed and shipped the event-driven order pipeline that "
            "processes 40 million events per day on AWS Lambda, EventBridge, "
            "and SQS. Cut p99 latency from 900ms to 120ms by replacing "
            "synchronous fan-out with a buffered consumer.",
            "Led the migration of the billing service from a monolith onto ECS "
            "Fargate. Wrote the Terraform modules, the blue/green deploy "
            "pipeline, and the rollback runbook the on-call rotation still "
            "follows.",
            "Introduced Apache Kafka as the backbone for cross-team data "
            "exchange: 14 topics, Avro contracts enforced by a schema registry, "
            "and consumer-lag alerting that catches a stalled consumer within "
            "two minutes.",
            "Owned the CI/CD platform for roughly 40 engineers. GitHub Actions "
            "runners on autoscaling EC2 plus a shared build cache brought mean "
            "build time from 11 minutes down to 4.",
            "Rebuilt the on-call rotation after a quarter averaging six pages a "
            "week. Wrote the postmortem process, ran 23 postmortems over 18 "
            "months, and tracked action items to completion; repeat incidents "
            "fell by about half.",
            "Led a team of five engineers through two performance cycles, "
            "including writing promotion packets for two of them.",
            "",
            "Software Engineer, Harbor Payments (2018 - 2021)",
            "Built and operated the settlement service handling roughly $2M in "
            "daily transaction volume. Designed the idempotency and retry model "
            "that made duplicate settlement runs safe.",
            "Migrated the primary datastore from MySQL to PostgreSQL with under "
            "four minutes of write downtime, using logical replication and a "
            "dual-write verification window.",
            "Added distributed tracing across seven services with OpenTelemetry, "
            "which turned a recurring 'the API is slow' complaint into a "
            "specific N+1 query in one downstream service.",
            "",
            "Software Engineer, Harbor Payments (2016 - 2018)",
            "Maintained internal REST APIs in Java and Python. Wrote the "
            "integration test suite that made the twice-weekly release train "
            "possible.",
            "Built the internal rate-limiting library that three teams adopted, "
            "after two outages traced back to a single client retrying "
            "aggressively without backoff.",
            "",
            "Backend Developer, Ridgeline Software (2015 - 2016)",
            "First engineering job. Wrote CRUD services in Java against an "
            "Oracle database for a warehouse client, and learned most of what I "
            "know about writing code other people have to read.",
            "Automated a manual weekly inventory reconciliation that had been "
            "taking an analyst most of a day, which is still the change I am "
            "proudest of relative to its size.",
        ],
    ),
    (
        "Selected work",
        [
            "Idempotency toolkit: a small library making retried writes safe "
            "across four services, built after a duplicate-settlement incident. "
            "Keyed on a client-supplied request id with a 24-hour dedup window "
            "in Redis, falling back to a database constraint when the cache is "
            "cold - because the cache is exactly what fails during an incident.",
            "Schema migration runbook: documented the expand/contract pattern "
            "the team now uses for every column change, after a rename shipped "
            "in a single deploy took the checkout service down for 11 minutes.",
            "Cost review: found that 30% of the AWS bill came from a logging "
            "sidecar shipping DEBUG output from every pod. Sampling brought it "
            "down without losing the traces we actually used.",
        ],
    ),
    (
        "Talks and writing",
        [
            "Internal tech talk: 'What our postmortems actually taught us', a "
            "review of 23 incidents and the three root causes that accounted "
            "for most of them.",
            "Wrote the team's service-design checklist covering SLOs, "
            "idempotency, backpressure, and rollback, now used at design "
            "review.",
        ],
    ),
    (
        "Skills",
        [
            "Languages: Go, Python, Java",
            "Infrastructure: AWS (Lambda, ECS, EventBridge, SQS, S3, EC2), "
            "Terraform, Docker, Kubernetes",
            "Data: PostgreSQL, MySQL, Redis, Apache Kafka",
            "Operations: Datadog, OpenTelemetry, GitHub Actions, incident "
            "response, SLOs and error budgets",
        ],
    ),
    (
        "Education",
        [
            "B.S. Computer Engineering, Gulf Coast Institute of Technology (2016)",
        ],
    ),
]


def write_docx(sections: list[tuple[str, list[str]]], path: Path) -> None:
    """Write a plain, single-column .docx.

    Heading levels are used rather than bare bold text so the file exercises
    the same style-aware path a résumé exported from Word would.
    """
    document = Document()
    for heading, paragraphs in sections:
        if heading:
            document.add_heading(heading, level=1)
        for text in paragraphs:
            document.add_paragraph(text)
    document.save(str(path))


def write_pdf(sections: list[tuple[str, list[str]]], path: Path) -> None:
    """Write a plain, single-column .pdf.

    Deliberately multi-page, matching the real file it replaces: a two-page
    PDF exercises page-boundary text extraction, which is where naive parsing
    tends to drop or duplicate content.
    """
    pdf = FPDF()
    pdf.set_margin(15.0)
    pdf.add_page()
    for heading, paragraphs in sections:
        if heading:
            pdf.set_font("Helvetica", "B", 13)
            pdf.multi_cell(0, 8, text=heading, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_font("Helvetica", size=11)
        for text in paragraphs:
            if not text:
                pdf.ln(3)
                continue
            pdf.multi_cell(0, 5.5, text=text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(3)
    path.write_bytes(bytes(pdf.output()))


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    docx_path = OUTPUT_DIR / "AI_Engineer_Resume.docx"
    pdf_path = OUTPUT_DIR / "Senior_SDE_Resume.pdf"

    write_docx(AI_ENGINEER, docx_path)
    write_pdf(SENIOR_SDE, pdf_path)

    for path in (docx_path, pdf_path):
        print(f"wrote {path.relative_to(OUTPUT_DIR.parents[1])} ({path.stat().st_size:,} bytes)")


if __name__ == "__main__":
    main()
